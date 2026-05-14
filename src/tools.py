import subprocess
import tempfile
import os
import base64
import mimetypes

import cv2
from langchain_core.messages import HumanMessage
from langchain_core.tools import tool

from config import SANDBOX_IMAGE, MAX_FRAMES
from models import vlm, llm
from ddgs import DDGS


@tool
def search_web(query: str) -> str:
    """Search the web for current information, recent events,
    or anything that requires up-to-date knowledge beyond
    your training data. Input should be a search query string."""
    with DDGS() as ddgs:
        results = ddgs.text(query, max_results=3)
        # Return early when the search yields nothing useful
        if not results:
            return "No results found."
        # Format each result into a readable title/url/summary block and join them
        return "\n".join(
            [
                f"Title: {r['title']}\nURL: {r['href']}\nSummary: {r['body']}"
                for r in results
            ]
        )


@tool
def execute_code(code: str) -> str:
    """Execute Python code safely in an isolated Docker container.
    Use this when the user asks to run code, perform calculations,
    or test a Python script. Input should be valid Python code."""
    tmp_path = None
    try:
        # Write the code to a temp file so Docker can mount it read-only into the container
        with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False) as tmp:
            tmp.write(code)
            tmp_path = tmp.name

        # Resource-constrained Docker run: no network, 128MB RAM cap, 0.5 CPU, 30s timeout
        result = subprocess.run(
            [
                "docker",
                "run",
                "--rm",
                "--network",
                "none",
                "--memory",
                "128m",
                "--cpus",
                "0.5",
                "-v",
                f"{tmp_path}:/app/code.py:ro",
                SANDBOX_IMAGE,
                "python",
                "/app/code.py",
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )

        # Distinguish a clean exit from a non-zero error exit
        if result.returncode == 0:
            return result.stdout or "Code executed successfully with no output."
        else:
            return f"Error:\n{result.stderr}"

    # Surface timeout as a user-readable message instead of a raw exception
    except subprocess.TimeoutExpired:
        return "Error: Code execution timed out after 30 seconds."
    # Catch Docker not found, permission errors, or any other unexpected failure
    except Exception as e:
        return f"Error: {e}"
    finally:
        # Always delete the temp file even if execution failed or timed out
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


@tool
def analyze_image(image_path: str, question: str = "What is in this image?") -> str:
    """Analyze an image using vision AI. Use this when the user
    provides an image path and wants to know what's in it,
    extract text from it, or ask questions about it.
    Input should be the path to the image file."""
    # Guard against a path that does not exist on disk
    if not os.path.exists(image_path):
        return f"Error: Image file not found at {image_path}"

    with open(image_path, "rb") as f:
        base64_image = base64.b64encode(f.read()).decode("utf-8")

    # Detect the actual MIME type from the file extension so PNG, GIF, and WEBP are labeled correctly
    mime_type = mimetypes.guess_type(image_path)[0] or "image/jpeg"

    message = HumanMessage(
        content=[
            {"type": "text", "text": question},
            {
                "type": "image_url",
                "image_url": {"url": f"data:{mime_type};base64,{base64_image}"},
            },
        ]
    )

    response = vlm.invoke([message])
    # Return a fallback string if the vision model returned an empty response
    return response.content if response.content else "No response from vision model."


@tool
def analyze_video(
    video_path: str,
    question: str = "What is happening in this video?",
) -> str:
    """Analyze a video using vision AI by sampling key frames.
    Use this when the user provides a video file path (.mp4, .avi, .mov, .mkv).
    Input should be the path to the video file."""
    # Guard against a path that does not exist on disk
    if not os.path.exists(video_path):
        return f"Error: Video file not found at {video_path}"

    cap = cv2.VideoCapture(video_path)
    # Guard against a file OpenCV cannot decode or open
    if not cap.isOpened():
        return f"Error: Could not open video file at {video_path}"

    try:
        # Compute a uniform frame interval so exactly MAX_FRAMES samples are spread across the video
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        interval = max(1, total_frames // MAX_FRAMES)
        frames_analyzed = 0
        content = [{"type": "text", "text": question}]

        # Walk through the video at the computed interval collecting sample frames
        for frame_idx in range(0, total_frames, interval):
            # Stop once the target number of frames has been collected
            if frames_analyzed >= MAX_FRAMES:
                break

            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
            ret, frame = cap.read()
            # Skip any frame that OpenCV failed to decode
            if not ret:
                continue

            # Resize and compress each frame to reduce the total payload sent to the VLM
            frame = cv2.resize(frame, (512, 512))
            _, buffer = cv2.imencode(".jpg", frame, [cv2.IMWRITE_JPEG_QUALITY, 70])
            base64_image = base64.b64encode(buffer.tobytes()).decode("utf-8")
            content.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"},
                }
            )
            frames_analyzed += 1
            print(f"Extracted frame {frames_analyzed}/{MAX_FRAMES}...", flush=True)

        # Nothing usable was extracted, report failure before calling the model
        if frames_analyzed == 0:
            return "No frames could be extracted from this video."

        print("Sending frames to vision model...", flush=True)
        response = vlm.invoke([HumanMessage(content=content)])
        return response.content or "No response from vision model."

    finally:
        cap.release()


@tool
def analyze_document(file_path: str, question: str = "Summarize this document") -> str:
    """Analyze a document file (PDF, Word, Excel, CSV).
    Use when user provides a .pdf, .docx, .xlsx, .xls, or .csv file path.
    Input should be the path to the document file."""

    # Guard against a path that does not exist on disk
    if not os.path.exists(file_path):
        return f"Error: File not found at {file_path}"

    ext = os.path.splitext(file_path)[1].lower()
    extracted_text = ""

    try:
        # Select the appropriate parser based on the file extension
        if ext == ".pdf":
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                # Accumulate text from every page in the document
                for page in pdf.pages:
                    text = page.extract_text()
                    # Skip pages that contain no extractable text
                    if text:
                        extracted_text += text + "\n"

        elif ext == ".docx":
            from docx import Document

            doc = Document(file_path)
            # Walk each paragraph and skip blank ones
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text += para.text + "\n"
            # Walk each table row and join cells with a pipe separator
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    extracted_text += row_text + "\n"

        elif ext in (".xlsx", ".xls"):
            import openpyxl

            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            # Iterate every sheet so multi-sheet workbooks are fully covered
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                extracted_text += f"Sheet: {sheet_name}\n"
                # Walk each data row and join non-None cells with a pipe separator
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(
                        str(cell) if cell is not None else "" for cell in row
                    )
                    # Skip rows that are entirely empty
                    if row_text.strip(" |"):
                        extracted_text += row_text + "\n"

        elif ext == ".csv":
            import pandas as pd

            df = pd.read_csv(file_path)
            extracted_text = df.to_string(index=False)

        else:
            return f"Error: Unsupported file type {ext}"

    # Catch library-level parsing errors and surface them as a readable message
    except Exception as e:
        return f"Error reading file: {e}"

    # Return early when the parser produced no usable content
    if not extracted_text.strip():
        return "Error: Could not extract any text from this file."

    # Truncate extracted text to stay within the model context window
    max_chars = 6000
    if len(extracted_text) > max_chars:
        extracted_text = extracted_text[:max_chars] + "\n... [truncated]"

    # Send extracted text alongside the user question to produce an answer
    response = llm.invoke([HumanMessage(content=f"""Here is the content of the document:

{extracted_text}

Question: {question}

Please answer the question based on the document content.""")])

    return response.content or "No response generated."


tools = [search_web, execute_code, analyze_image, analyze_video, analyze_document]
